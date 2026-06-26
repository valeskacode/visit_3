# -*- coding: utf-8 -*-
"""
utils/helpers.py
Funciones compartidas por toda la app: estilos, lectura de Excel,
autoguardado/recuperación de avance, historial y generación de reportes
(Word y PDF).
"""
import io
import json
import os
import re
import shutil
from datetime import datetime

import pandas as pd
import streamlit as st

try:
    from zoneinfo import ZoneInfo
    TZ_PERU = ZoneInfo("America/Lima")
except Exception:  # pragma: no cover - fallback si no hay tzdata
    TZ_PERU = None


def ahora_peru() -> datetime:
    """Devuelve la fecha/hora ACTUAL de Perú (America/Lima, UTC-5),
    sin importar en qué servidor/huso horario corra Streamlit Cloud.
    Usar SIEMPRE esta función en vez de datetime.now() para que la
    hora/fecha de visita se autocomplete correctamente en Perú."""
    if TZ_PERU is not None:
        return datetime.now(TZ_PERU)
    return datetime.now()


# --------------------------------------------------------------------------
# RUTAS DE DATOS LOCALES (persisten mientras el servidor no se reinicie)
# --------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, "data")
DRAFTS_DIR = os.path.join(DATA_DIR, "drafts")
FOTOS_DIR = os.path.join(DRAFTS_DIR, "fotos")
HISTORIAL_PATH = os.path.join(DATA_DIR, "historial_visitas.xlsx")

# --------------------------------------------------------------------------
# 📁 CARPETA DE ALMACENAMIENTO DE REPORTES (Word/PDF) — EDITAR AQUÍ 👇
# --------------------------------------------------------------------------
# Cambia esta ruta a la carpeta donde quieres que se guarden TODOS los
# reportes generados (Word y PDF) cada vez que alguien presiona
# "Generar Word" / "Generar PDF" en la vista de Reporte.
#
# Para que terminen sincronizados en OneDrive (sin necesitar credenciales
# de Microsoft Graph API, que requieren registro de app/permisos), la
# forma más simple y confiable es apuntar esta ruta a una carpeta que YA
# esté dentro de tu OneDrive sincronizado en este equipo/servidor:
#   Windows:    REPORTES_DIR = r"C:\Users\TU_USUARIO\OneDrive\Auditoria\Visitas"
#   Linux/macOS (con OneDrive vía rclone o cliente "onedrive"):
#               REPORTES_DIR = "/home/tu_usuario/OneDrive/Auditoria/Visitas"
# También puedes definir la variable de entorno VISITA_APP_REPORTES_DIR
# al desplegar la app, sin tocar el código.
REPORTES_DIR = os.environ.get(
    "VISITA_APP_REPORTES_DIR", os.path.join(DATA_DIR, "reportes_generados")
)

for _d in (DATA_DIR, DRAFTS_DIR, FOTOS_DIR, REPORTES_DIR):
    os.makedirs(_d, exist_ok=True)

EXCEL_SHEET_NAME = "MUESTRA_FINAL"

EXCEL_COLUMNS = [
    "RECNO", "PEPAIS", "PETDOC", "DOCPEN", "CODCLI", "BCEMP", "BCSUC", "BCMDA",
    "BCPAP", "BCCTA", "BCOPER", "BCSBOP", "BCTOP", "BCMOD", "CODCRE", "REGION",
    "ZONA", "AGENCIA", "CLIENTE", "DIRECCION_DOM", "DISTRITO_DOM",
    "PROVINCIA_DOM", "DEPARTAMENTO_DOM", "DIRECCION_NEG", "DISTRITO_NEG",
    "PROVINCIA_NEG", "DEPARTAMENTO_NEG", "ACTIVIDAD_ECON", "ANALISTA",
    "PRODUCTO_CAJA", "SALDO_MN", "SALDO_VIGE", "SALDO_REFI", "SALDO_VENC",
    "SALDO_JUDI", "MORA_CONT", "TIPO_SBS", "FECDES", "IMPDESEMB_MN",
    "COD_MODULO", "MODULO", "COD_TIPO_OPERACION", "TIPO_OPERACION",
    "ANALISTA_EVAL", "USUARIO_APROB", "USUARIO_DESEM", "FECHA_EVAL",
    "DIAS_ATRASO", "ESTADO_CREDITO", "ATRANT_1M", "ATRANT_2M", "ATRANT_3M",
    "ATRANT_4M", "ATRANT_5M", "ATRANT_6M", "TIPO_SOLI", "NUMERO_CUOTAS",
    "CUOTAS_PAGADAS", "TIPO", "SEGMENTACION_MYPE", "CATEG_RESULTANTE",
    "CATEG_RESULTANTE_SINALIN", "CUENTA_AVAL", "FECHA_UTLPAGO", "UAI_IND",
    "ESTRATO", "TIPO_EXPEDIENTE",
]

# --------------------------------------------------------------------------
# 📌 NUEVO — COLUMNAS OPCIONALES PARA LA TARJETA "Cliente encontrado"
# --------------------------------------------------------------------------
# La tarjeta de búsqueda (pantalla_busqueda → render_cliente_encontrado en
# app.py) ahora muestra teléfono, correo y límite de crédito, tal como el
# mockup. Tu Excel actual (columnas de arriba) NO trae esos tres datos, así
# que por ahora se muestran como "No registrado" / "No disponible".
#
# EDITAR AQUÍ 👇 si tu base de datos sí tiene esas columnas: escribe el
# nombre EXACTO de la columna (en mayúsculas, como queda tras leer el
# Excel) en vez de "TELEFONO" / "EMAIL" / "LIMITE_CREDITO_MN", y usa ese
# mismo nombre en `render_cliente_encontrado()` dentro de app.py (sección
# marcada con 🔧 AQUÍ en ese archivo).
COLUMNAS_OPCIONALES_CONTACTO = [
    "TELEFONO",           # Ej: número de celular del cliente
    "EMAIL",               # Ej: correo de contacto
    "LIMITE_CREDITO_MN",   # Ej: línea de crédito aprobada (no es el saldo actual)
]
EXCEL_COLUMNS = EXCEL_COLUMNS + [
    c for c in COLUMNAS_OPCIONALES_CONTACTO if c not in EXCEL_COLUMNS
]

CLIENTE_VISITADO_OPCIONES = [
    "1. Cliente con actividad laboral y/o económica vigente",
    "2. Cliente con situación desmejorada",
    "3. Cliente ya no labora y/o no realiza la actividad económica",
    "4. Cliente no ubicado",
]

CRITERIOS_DEF = {
    "Indicio de dolo o fraude en la evaluación de créditos": [
        "Documentos con enmendaduras",
        "Documentos con datos inconsistentes",
        "Documentos sin datos del cliente",
        "Documentos sin firmas o que no coinciden",
        "Documentos duplicados en más de un cliente",
    ],
    "Evaluaciones deficientes o con sustento insuficiente": [
        "No se evidenció sustento de actividad económica",
        "No se evidenció sustento de ingresos",
        "No se evidenció sustento de activos representativos",
        "Se omitió al cónyuge",
    ],
    "Créditos reprogramados y refinanciados": [
        "Reprogramado",
        "Refinanciado",
    ],
    "Clientes con créditos con calificación diferente a normal a la fecha de revisión": [
        "Calificación diferente a normal",
    ],
}


# --------------------------------------------------------------------------
# ESTILOS
# --------------------------------------------------------------------------
def load_css(path):
    """Inyecta un archivo CSS dentro de la app de Streamlit."""
    full_path = os.path.join(BASE_DIR, path) if not os.path.isabs(path) else path
    try:
        with open(full_path, "r", encoding="utf-8") as f:
            css = f.read()
        st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)
    except FileNotFoundError:
        st.warning(f"No se encontró el archivo de estilos: {full_path}")


# --------------------------------------------------------------------------
# HELPERS DE DATOS
# --------------------------------------------------------------------------
def safe_str(v, default=""):
    if v is None:
        return default
    try:
        if pd.isna(v):
            return default
    except Exception:
        pass
    s = str(v).strip()
    return default if s.lower() in ("nan", "none") else s


def safe_float(v, default=0.0):
    try:
        f = float(str(v).replace(",", "").strip())
        if pd.isna(f):
            return default
        return f
    except Exception:
        return default


def fmt_money(v):
    return f"S/ {safe_float(v):,.2f}"


def slug(texto):
    """Convierte un texto en algo seguro para usar como nombre de archivo
    o clave de widget, quitando tildes correctamente en vez de dejarlas
    como guiones bajos sueltos (ej. 'Pérez' -> 'Perez', no 'P_rez')."""
    import unicodedata
    texto = safe_str(texto, "sin_dato")
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r"[^A-Za-z0-9_\-]+", "_", texto)
    texto = re.sub(r"_+", "_", texto)
    return texto.strip("_") or "sin_dato"


# --------------------------------------------------------------------------
# 📌 NUEVO — HELPERS PARA LA TARJETA "Cliente encontrado" Y LA LISTA DE
# "Clientes similares" de la pantalla de Búsqueda (igualando el mockup).
# --------------------------------------------------------------------------
def solo_digitos(texto):
    """Deja solo los dígitos de un texto. Útil para comparar DNIs aunque
    el usuario escriba espacios, puntos o el nombre junto al número."""
    return re.sub(r"\D", "", safe_str(texto))


def iniciales(nombre):
    """Devuelve hasta 2 iniciales en mayúscula para el avatar circular
    (ej. 'Juan Pérez García' -> 'JP')."""
    partes = [p for p in safe_str(nombre).split() if p]
    if not partes:
        return "?"
    if len(partes) == 1:
        return partes[0][:2].upper()
    return (partes[0][0] + partes[1][0]).upper()


def clase_calificacion(calif):
    """Clase de color para el 'chip' de calificación (A=verde, B=ámbar,
    cualquier otra cosa=rojo). Ajusta este mapeo si tu escala SBS usa
    otras letras o un formato distinto."""
    c = safe_str(calif).strip().upper()
    if c.startswith("A"):
        return "chip-calif-ok"
    if c.startswith("B"):
        return "chip-calif-warn"
    if c:
        return "chip-calif-bad"
    return "chip-calif-na"


def clientes_similares(df, fila_actual, max_resultados=3):
    """Busca otros clientes con nombre o DNI parecido al de `fila_actual`,
    para ayudar al usuario a distinguir homónimos o detectar posibles
    duplicados (ver criterio "Documentos duplicados en más de un cliente"
    en CRITERIOS_DEF). No es una librería de fuzzy-matching: es una
    heurística simple a propósito, para no añadir dependencias nuevas.

    Se considera "similar" un registro que:
      - comparte 2 o más palabras del nombre completo, o
      - comparte los primeros 6 dígitos del DNI,
    y que NO es exactamente el mismo registro (mismo DNI Y mismo nombre).
    """
    nombre_actual = safe_str(fila_actual.get("CLIENTE")).strip().lower()
    dni_actual = solo_digitos(fila_actual.get("DOCPEN"))
    palabras_actual = set(nombre_actual.split())

    candidatos = []
    for _, row in df.iterrows():
        nombre_row = safe_str(row.get("CLIENTE")).strip().lower()
        dni_row = solo_digitos(row.get("DOCPEN"))
        if dni_row == dni_actual and nombre_row == nombre_actual:
            continue  # es el mismo cliente, no un "similar"
        palabras_row = set(nombre_row.split())
        comparten_nombre = len(palabras_actual & palabras_row) >= 2
        comparten_dni = len(dni_actual) >= 6 and dni_row[:6] == dni_actual[:6]
        if comparten_nombre or comparten_dni:
            candidatos.append(row)
        if len(candidatos) >= max_resultados:
            break
    return pd.DataFrame(candidatos) if candidatos else pd.DataFrame(columns=df.columns)


# --------------------------------------------------------------------------
# LECTURA DEL EXCEL (hoja MUESTRA_FINAL) — con caché para que cargue rápido
# tanto en PC como en celular (el procesamiento ocurre en el servidor, no
# en el dispositivo, así que cachear el resultado evita reprocesar el
# archivo en cada clic).
# --------------------------------------------------------------------------
@st.cache_data(show_spinner="Procesando archivo Excel...")
def cargar_excel(file_bytes: bytes):
    bio = io.BytesIO(file_bytes)
    hoja_usada = EXCEL_SHEET_NAME
    try:
        df = pd.read_excel(bio, sheet_name=EXCEL_SHEET_NAME, dtype=str)
    except ValueError:
        bio.seek(0)
        xls = pd.ExcelFile(bio)
        hoja_usada = xls.sheet_names[0]
        df = pd.read_excel(bio, sheet_name=hoja_usada, dtype=str)

    df.columns = [str(c).strip().upper() for c in df.columns]
    # Compatibilidad con archivos antiguos que usaban "PENDOC" en vez de "DOCPEN"
    if "PENDOC" in df.columns and "DOCPEN" not in df.columns:
        df = df.rename(columns={"PENDOC": "DOCPEN"})
    df = df.fillna("")
    faltantes = [c for c in EXCEL_COLUMNS if c not in df.columns]
    return df, hoja_usada, faltantes


# --------------------------------------------------------------------------
# AUTOGUARDADO / RECUPERACIÓN DE AVANCE
# --------------------------------------------------------------------------
INGRESOS_KEYS = [
    "ingreso_principal", "otros_ingresos",
    "op_alquiler", "op_servicios", "op_transporte", "op_mercaderia", "op_publicidad", "op_otros",
    "fam_alimentacion", "fam_vivienda", "fam_servicios", "fam_educacion", "fam_salud", "fam_otros",
]


def _draft_path(usuario, dni):
    return os.path.join(DRAFTS_DIR, f"{slug(usuario)}__{slug(dni)}.json")


def hay_borrador(usuario, dni):
    return os.path.exists(_draft_path(usuario, dni))


def guardar_borrador(usuario, dni, cliente_actual):
    """Guarda el avance actual de la sesión a disco (foto incluida).

    Los checkboxes de criterios y los montos de ingresos/gastos se guardan
    "planos" (mismas claves que usan sus widgets) para poder restaurarlos
    directamente en session_state antes de que esos widgets se vuelvan a
    dibujar — así Streamlit los muestra ya marcados/llenados al recuperar.
    """
    if not usuario or not dni:
        return
    data = {"cliente_actual": cliente_actual, "guardado_en": ahora_peru().isoformat()}

    data["criterios"] = {k: v for k, v in st.session_state.items() if k.startswith("chk_")}
    data["calif_revision"] = st.session_state.get("calif_revision", "")
    data["ingresos_gastos"] = {k: st.session_state.get(k, 0.0) for k in INGRESOS_KEYS}
    data["garantias"] = st.session_state.get("garantias", [])
    data["rcc"] = st.session_state.get("rcc", [])
    data["cliente_visitado"] = st.session_state.get("cliente_visitado", "")

    visitas_serializables = {}
    for clave, v in st.session_state.get("visitas", {}).items():
        v = dict(v)
        foto_bytes = v.pop("foto_bytes", None)
        if foto_bytes:
            foto_path = os.path.join(FOTOS_DIR, f"{slug(usuario)}__{slug(dni)}__{clave}.jpg")
            with open(foto_path, "wb") as f:
                f.write(foto_bytes)
            v["foto_path"] = foto_path
        visitas_serializables[clave] = v
    data["visitas"] = visitas_serializables

    with open(_draft_path(usuario, dni), "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, default=str)


def cargar_borrador(usuario, dni):
    """Carga un avance guardado previamente de vuelta a session_state."""
    path = _draft_path(usuario, dni)
    if not os.path.exists(path):
        return False
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    for k, v in data.get("criterios", {}).items():
        st.session_state[k] = v
    st.session_state["calif_revision"] = data.get("calif_revision", "")

    for k, v in data.get("ingresos_gastos", {}).items():
        st.session_state[k] = v

    st.session_state["garantias"] = data.get("garantias", [])
    st.session_state["rcc"] = data.get("rcc", [])
    st.session_state["cliente_visitado"] = data.get("cliente_visitado", "")

    visitas = data.get("visitas", {})
    for clave, v in visitas.items():
        foto_path = v.pop("foto_path", None)
        if foto_path and os.path.exists(foto_path):
            with open(foto_path, "rb") as imgf:
                v["foto_bytes"] = imgf.read()
    st.session_state["visitas"] = visitas
    return True


def borrar_borrador(usuario, dni):
    path = _draft_path(usuario, dni)
    if os.path.exists(path):
        os.remove(path)
    for clave in ("domicilio", "negocio", "aval"):
        foto_path = os.path.join(FOTOS_DIR, f"{slug(usuario)}__{slug(dni)}__{clave}.jpg")
        if os.path.exists(foto_path):
            os.remove(foto_path)


# --------------------------------------------------------------------------
# HISTORIAL (registro de quién generó qué informe y cuándo)
# --------------------------------------------------------------------------
HISTORIAL_COLUMNS = [
    "Usuario_Auditor", "Agencia", "Cliente", "DNI", "Cuenta",
    "N_Visita_Agencia", "N_Visita_General",
    "ClienteVisitado", "Fecha", "Hora",
    "TipoArchivo", "NombreArchivo", "RutaGuardado", "CriteriosSeleccionados",
]


def leer_historial():
    if not os.path.exists(HISTORIAL_PATH):
        return pd.DataFrame(columns=HISTORIAL_COLUMNS)
    return pd.read_excel(HISTORIAL_PATH)


def _contar_visitas_previas(agencia):
    """Cuenta cuántas visitas ya hay registradas para una agencia dada y
    en total, para asignar el numerador de la visita siguiente."""
    hist = leer_historial()
    n_general = len(hist)
    if "Agencia" in hist.columns and agencia:
        n_agencia = int((hist["Agencia"].astype(str).str.strip() == str(agencia).strip()).sum())
    else:
        n_agencia = 0
    return n_agencia, n_general


def registrar_historial(usuario, cliente_actual, tipo_archivo, nombre_archivo,
                         criterios_texto, cliente_visitado="", ruta_guardado=""):
    """Agrega una fila al historial general (data/historial_visitas.xlsx).

    Calcula automáticamente N_Visita_Agencia (numerador dentro de esa
    agencia) y N_Visita_General (numerador global, todas las agencias).
    """
    import openpyxl

    agencia = safe_str(cliente_actual.get("AGENCIA"))
    n_agencia_prev, n_general_prev = _contar_visitas_previas(agencia)

    ahora = ahora_peru()
    fila = [
        usuario, agencia,
        safe_str(cliente_actual.get("CLIENTE")),
        safe_str(cliente_actual.get("DOCPEN")),
        safe_str(cliente_actual.get("BCCTA")),
        n_agencia_prev + 1, n_general_prev + 1,
        cliente_visitado,
        ahora.strftime("%d/%m/%Y"), ahora.strftime("%H:%M:%S"),
        tipo_archivo, nombre_archivo, ruta_guardado, criterios_texto,
    ]
    if os.path.exists(HISTORIAL_PATH):
        wb = openpyxl.load_workbook(HISTORIAL_PATH)
        ws = wb.active
        encabezado_actual = [c.value for c in ws[1]]
        if encabezado_actual != HISTORIAL_COLUMNS:
            # Migra archivos de historial de una versión anterior de la app
            # sin perder las filas ya registradas.
            ws.delete_rows(1)
            ws.insert_rows(1)
            for i, col in enumerate(HISTORIAL_COLUMNS, start=1):
                ws.cell(row=1, column=i, value=col)
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Historial"
        ws.append(HISTORIAL_COLUMNS)
    ws.append(fila)
    wb.save(HISTORIAL_PATH)
    return n_agencia_prev + 1, n_general_prev + 1


def reporte_consolidado_por_agencia():
    """Cuántos clientes distintos se visitó por agencia y cuántos
    reportes en total, a partir de todo lo guardado en el historial."""
    hist = leer_historial()
    if hist.empty or "Agencia" not in hist.columns:
        return pd.DataFrame(columns=["Agencia", "Clientes_Visitados", "Reportes_Generados", "Ultima_Visita"])
    resumen = (
        hist.groupby("Agencia")
        .agg(Clientes_Visitados=("DNI", "nunique"),
             Reportes_Generados=("NombreArchivo", "count"),
             Ultima_Visita=("Fecha", "max"))
        .reset_index()
        .sort_values("Clientes_Visitados", ascending=False)
    )
    return resumen


def reporte_consolidado_por_cliente(agencia=None):
    """Detalle por cliente: cuántas visitas/reportes tiene cada uno,
    opcionalmente filtrado por agencia."""
    hist = leer_historial()
    if hist.empty:
        return pd.DataFrame(columns=["Agencia", "Cliente", "DNI", "Cuenta", "Reportes_Generados", "Ultima_Visita"])
    if agencia and "Agencia" in hist.columns:
        hist = hist[hist["Agencia"].astype(str).str.strip() == str(agencia).strip()]
    resumen = (
        hist.groupby(["Agencia", "Cliente", "DNI", "Cuenta"], dropna=False)
        .agg(Reportes_Generados=("NombreArchivo", "count"), Ultima_Visita=("Fecha", "max"))
        .reset_index()
        .sort_values("Reportes_Generados", ascending=False)
    )
    return resumen


def guardar_reporte_en_carpeta(nombre_archivo, contenido_bytes):
    """Guarda una copia física del reporte en REPORTES_DIR (ver
    definición arriba de esta carpeta — puede apuntar a una carpeta
    sincronizada con OneDrive). Devuelve la ruta final, o "" si falló."""
    try:
        destino = os.path.join(REPORTES_DIR, nombre_archivo)
        with open(destino, "wb") as f:
            f.write(contenido_bytes)
        return destino
    except Exception:
        return ""


# --------------------------------------------------------------------------
# CÁLCULOS DE EVALUACIÓN
# --------------------------------------------------------------------------
def calcular_resultado(ing):
    total_ingresos = safe_float(ing.get("ingreso_principal")) + safe_float(ing.get("otros_ingresos"))
    gastos_operativos = sum(safe_float(ing.get(k)) for k in [
        "op_alquiler", "op_servicios", "op_transporte", "op_mercaderia", "op_publicidad", "op_otros",
    ])
    gastos_familiares = sum(safe_float(ing.get(k)) for k in [
        "fam_alimentacion", "fam_vivienda", "fam_servicios", "fam_educacion", "fam_salud", "fam_otros",
    ])
    total_gastos = gastos_operativos + gastos_familiares
    utilidad_neta = total_ingresos - total_gastos
    margen = (utilidad_neta / total_ingresos * 100) if total_ingresos else 0.0
    return {
        "total_ingresos": total_ingresos,
        "gastos_operativos": gastos_operativos,
        "gastos_familiares": gastos_familiares,
        "total_gastos": total_gastos,
        "utilidad_neta": utilidad_neta,
        "margen": margen,
    }


def criterios_seleccionados_lista(criterios, calif_revision):
    seleccionados = []
    for categoria, items in CRITERIOS_DEF.items():
        for item in items:
            key = f"chk_{slug(categoria)}_{slug(item)}"
            if criterios.get(key):
                if item == "Calificación diferente a normal" and calif_revision:
                    seleccionados.append(f"{item} ({calif_revision})")
                else:
                    seleccionados.append(item)
    return seleccionados


# --------------------------------------------------------------------------
# GENERACIÓN DE REPORTE — WORD
# --------------------------------------------------------------------------
def generar_word(cliente, criterios_txt, ingresos_calc, ingresos_raw, visitas, garantias, rcc, usuario, cliente_visitado=""):
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    AZUL = "1B3A5C"

    def add_heading(doc, text, size=13):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(AZUL)
        return p

    def add_kv_table(doc, pairs, cols=2):
        table = doc.add_table(rows=0, cols=cols * 2)
        table.style = "Light Grid Accent 1"
        row = None
        for i, (k, v) in enumerate(pairs):
            if i % cols == 0:
                row = table.add_row().cells
            c = (i % cols) * 2
            row[c].text = str(k)
            row[c + 1].text = str(v) if v not in (None, "") else "-"
        return table

    doc = Document()
    doc.add_heading("VISITA A CLIENTES DE PEQUEÑA EMPRESA", level=0)
    p = doc.add_paragraph("CMAC Caja Arequipa — Unidad de Auditoría Interna")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Auditor: {usuario}  ·  Fecha de visita: {ahora_peru().strftime('%d/%m/%Y %H:%M')} (hora Perú)")

    if criterios_txt:
        add_heading(doc, "0. Criterio para la visita")
        for c in criterios_txt:
            doc.add_paragraph("• " + c)

    add_heading(doc, "I. Datos del cliente y crédito")
    add_kv_table(doc, [
        ("Agencia", safe_str(cliente.get("AGENCIA"))),
        ("DNI/LE Titular", safe_str(cliente.get("DOCPEN"))),
        ("Titular", safe_str(cliente.get("CLIENTE"))),
        ("Cuenta cliente", safe_str(cliente.get("BCCTA"))),
        ("N° de operación", safe_str(cliente.get("BCOPER"))),
        ("Módulo", safe_str(cliente.get("MODULO"))),
        ("Analista vigente", safe_str(cliente.get("ANALISTA"))),
        ("Analista evaluador", safe_str(cliente.get("ANALISTA_EVAL"))),
        ("Auditor (visita)", usuario),
        ("Importe", fmt_money(cliente.get("IMPDESEMB_MN"))),
        ("Saldo capital", fmt_money(cliente.get("SALDO_MN"))),
        ("Tipo de crédito", safe_str(cliente.get("PRODUCTO_CAJA"))),
        ("Tipo SBS", safe_str(cliente.get("TIPO_SBS"))),
        ("Calificación", safe_str(cliente.get("CATEG_RESULTANTE"))),
        ("Rubro", safe_str(cliente.get("ACTIVIDAD_ECON"))),
        ("Último pago", safe_str(cliente.get("FECHA_UTLPAGO"))),
        ("Resultado de la visita / Cliente visitado", cliente_visitado or "-"),
    ])

    add_heading(doc, "II. Ingresos y gastos")
    add_kv_table(doc, [
        ("Ingreso principal", fmt_money(ingresos_raw.get("ingreso_principal"))),
        ("Otros ingresos", fmt_money(ingresos_raw.get("otros_ingresos"))),
        ("Total ingresos", fmt_money(ingresos_calc["total_ingresos"])),
        ("Gastos operativos", fmt_money(ingresos_calc["gastos_operativos"])),
        ("Gastos familiares", fmt_money(ingresos_calc["gastos_familiares"])),
        ("Total gastos", fmt_money(ingresos_calc["total_gastos"])),
        ("Utilidad neta", fmt_money(ingresos_calc["utilidad_neta"])),
        ("Margen", f"{ingresos_calc['margen']:.1f}%"),
    ])

    for clave, titulo in [("negocio", "III. Visita al negocio (dirección del negocio)"),
                           ("laboral", "IV. Visita al centro laboral"),
                           ("aval", "V. Visita al aval"),
                           ("domicilio", "VI. Visita al domicilio")]:
        d = visitas.get(clave)
        add_heading(doc, titulo)
        if d:
            add_kv_table(doc, [
                ("Dirección", d.get("direccion", "-")),
                ("Distrito", d.get("distrito", "-")),
                ("Provincia", d.get("provincia", "-")),
                ("Departamento", d.get("departamento", "-")),
                ("Referencia", d.get("referencia", "-")),
                ("Fecha de visita", d.get("fecha", "-")),
                ("Hora de visita", d.get("hora", "-")),
                ("Entrevista con", d.get("entrevista_con", "-")),
                ("Comentarios", d.get("comentarios", "-")),
                ("GPS", f"{d.get('lat')}, {d.get('lon')}" if d.get("lat") else "No capturada"),
            ])
            if d.get("foto_bytes"):
                doc.add_picture(io.BytesIO(d["foto_bytes"]), width=Cm(8))
        else:
            doc.add_paragraph("⚠ No se registró visita de verificación para esta sección.")

    if garantias:
        add_heading(doc, "VI. Garantías")
        for g in garantias:
            add_kv_table(doc, list(g.items()))

    if rcc:
        add_heading(doc, "VII. Deuda RCC")
        for r in rcc:
            add_kv_table(doc, list(r.items()))

    add_heading(doc, "Conformidad")
    add_kv_table(doc, [
        ("Hecho por (Auditor)", usuario), ("Fecha", ahora_peru().strftime("%d/%m/%Y")),
        ("Revisado por", ""), ("Fecha", ""),
    ])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------
# GENERACIÓN DE REPORTE — PDF (independiente de Word, usando reportlab)
# --------------------------------------------------------------------------
def generar_pdf(cliente, criterios_txt, ingresos_calc, ingresos_raw, visitas, garantias, rcc, usuario, cliente_visitado=""):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    AZUL = colors.HexColor("#1B3A5C")
    ROJO = colors.HexColor("#C8102E")

    buf = io.BytesIO()
    docpdf = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1c", parent=styles["Heading1"], textColor=AZUL, fontSize=15)
    h2 = ParagraphStyle("h2c", parent=styles["Heading2"], textColor=AZUL, fontSize=12, spaceBefore=10)
    normal = styles["Normal"]

    elems = [
        Paragraph("VISITA A CLIENTES DE PEQUEÑA EMPRESA", h1),
        Paragraph("CMAC Caja Arequipa — Unidad de Auditoría Interna", normal),
        Paragraph(f"Auditor: {usuario} · Fecha de visita: {ahora_peru().strftime('%d/%m/%Y %H:%M')} (hora Perú)", normal),
        Spacer(1, 10),
    ]

    def tabla_kv(pairs):
        data = [[k, str(v) if v not in (None, "") else "-"] for k, v in pairs]
        t = Table(data, colWidths=[6 * cm, 9 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f0f3")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        return t

    if criterios_txt:
        elems.append(Paragraph("0. Criterio para la visita", h2))
        for c in criterios_txt:
            elems.append(Paragraph("• " + c, normal))

    elems.append(Paragraph("I. Datos del cliente y crédito", h2))
    elems.append(tabla_kv([
        ("Agencia", safe_str(cliente.get("AGENCIA"))),
        ("DNI/LE Titular", safe_str(cliente.get("DOCPEN"))),
        ("Titular", safe_str(cliente.get("CLIENTE"))),
        ("Cuenta cliente", safe_str(cliente.get("BCCTA"))),
        ("N° de operación", safe_str(cliente.get("BCOPER"))),
        ("Módulo", safe_str(cliente.get("MODULO"))),
        ("Analista vigente", safe_str(cliente.get("ANALISTA"))),
        ("Analista evaluador", safe_str(cliente.get("ANALISTA_EVAL"))),
        ("Auditor (visita)", usuario),
        ("Importe", fmt_money(cliente.get("IMPDESEMB_MN"))),
        ("Saldo capital", fmt_money(cliente.get("SALDO_MN"))),
        ("Tipo de crédito", safe_str(cliente.get("PRODUCTO_CAJA"))),
        ("Calificación", safe_str(cliente.get("CATEG_RESULTANTE"))),
        ("Resultado de la visita / Cliente visitado", cliente_visitado or "-"),
    ]))

    elems.append(Paragraph("II. Ingresos y gastos", h2))
    elems.append(tabla_kv([
        ("Total ingresos", fmt_money(ingresos_calc["total_ingresos"])),
        ("Gastos operativos", fmt_money(ingresos_calc["gastos_operativos"])),
        ("Gastos familiares", fmt_money(ingresos_calc["gastos_familiares"])),
        ("Total gastos", fmt_money(ingresos_calc["total_gastos"])),
        ("Utilidad neta", fmt_money(ingresos_calc["utilidad_neta"])),
        ("Margen", f"{ingresos_calc['margen']:.1f}%"),
    ]))

    for clave, titulo in [("negocio", "III. Visita al negocio (dirección del negocio)"),
                           ("laboral", "IV. Visita al centro laboral"),
                           ("aval", "V. Visita al aval"),
                           ("domicilio", "VI. Visita al domicilio")]:
        d = visitas.get(clave)
        elems.append(Paragraph(titulo, h2))
        if d:
            elems.append(tabla_kv([
                ("Dirección", d.get("direccion", "-")),
                ("Distrito", d.get("distrito", "-")),
                ("Fecha de visita", d.get("fecha", "-")),
                ("Hora de visita", d.get("hora", "-")),
                ("Entrevista con", d.get("entrevista_con", "-")),
                ("Comentarios", d.get("comentarios", "-")),
                ("GPS", f"{d.get('lat')}, {d.get('lon')}" if d.get("lat") else "No capturada"),
            ]))
            if d.get("foto_bytes"):
                try:
                    img = RLImage(io.BytesIO(d["foto_bytes"]), width=8 * cm, height=6 * cm)
                    elems.append(img)
                except Exception:
                    pass
        else:
            elems.append(Paragraph("No se registró visita de verificación para esta sección.", normal))
        elems.append(Spacer(1, 6))

    if garantias:
        elems.append(Paragraph("VI. Garantías", h2))
        for g in garantias:
            elems.append(tabla_kv(list(g.items())))

    if rcc:
        elems.append(Paragraph("VII. Deuda RCC", h2))
        for r in rcc:
            elems.append(tabla_kv(list(r.items())))

    elems.append(Paragraph("Conformidad", h2))
    elems.append(tabla_kv([
        ("Hecho por (Auditor)", usuario),
        ("Fecha", ahora_peru().strftime("%d/%m/%Y")),
    ]))

    docpdf.build(elems)
    buf.seek(0)
    return buf
